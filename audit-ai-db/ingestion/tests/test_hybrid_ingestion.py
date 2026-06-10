from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from ingestion.legacy.gemini_ingestion import run_gemini_ingestion as legacy_gemini_ingestion
from ingestion.advanced_markdown import detect_broken_table, repair_complex_tables
from ingestion.chunker import create_chunks
from ingestion.gemini.chunker import _chunk_records_from_json
from ingestion.gemini.reader import GeminiTokenUsage
from ingestion.legacy.local_ingestion import run_ingestion as legacy_local_ingestion
from ingestion.metadata_extractor import enrich_metadata_from_markdown, prepare_metadata
from ingestion.run_folder import run_folder_ingestion
from ingestion.hybrid.pipeline import run_hybrid_ingestion
from ingestion.hybrid.strategies import (
    PdfPageStats,
    choose_ingestion_strategy,
    select_pdf_strategy,
)
from ingestion.models import ChunkRecord, FileInfo, IngestionError, SummaryResult
from ingestion.models import StructuredSection
from ingestion.hybrid.chunkers import prepare_local_chunks
from ingestion.run_ingestion import run_ingestion


class HybridIngestionTests(unittest.TestCase):
    def test_auto_strategy_selects_local_for_docx(self) -> None:
        file_info = FileInfo(
            file_path=Path("policy.docx"),
            file_name="policy.docx",
            file_extension=".docx",
            file_size=100,
            file_type="docx",
        )

        decision = choose_ingestion_strategy(file_info, "auto")

        self.assertEqual(decision.selected_strategy, "local")
        self.assertEqual(decision.reason, "docx_text_first")

    def test_pdf_stats_select_gemini_for_image_heavy_low_text_pdf(self) -> None:
        stats = PdfPageStats(
            pages=10,
            total_text_chars=200,
            pages_with_images=10,
            low_text_pages=10,
            low_text_image_pages=10,
        )

        strategy, reason = select_pdf_strategy(stats)

        self.assertEqual(strategy, "gemini")
        self.assertIn("pdf_", reason)

    def test_invalid_strategy_raises_ingestion_error(self) -> None:
        file_info = FileInfo(
            file_path=Path("policy.docx"),
            file_name="policy.docx",
            file_extension=".docx",
            file_size=100,
            file_type="docx",
        )

        with self.assertRaises(IngestionError):
            choose_ingestion_strategy(file_info, "unknown")

    def test_local_no_db_hybrid_ingestion_chunks_docx(self) -> None:
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx is not installed")

        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "hybrid_policy.docx"
            document = Document()
            document.add_heading("資料共享管理辦法", level=1)
            document.add_paragraph(
                "第一條 客戶資料共享應載明於客戶已簽署之契據文件及官網個資應告知事項。"
            )
            document.add_paragraph(
                "第二條 涉及負面資訊或風險類資料時，應確認必要查證程序與保護措施。"
            )
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "項目"
            table.cell(0, 1).text = "要求"
            table.cell(1, 0).text = "資料共享"
            table.cell(1, 1).text = "應保留紀錄"
            document.save(docx_path)
            output_dir = Path(temp_dir) / "processed"

            result = run_hybrid_ingestion(
                {
                    "file_path": str(docx_path),
                    "internal_code": "HYBRID-TEST-001",
                    "document_type": "internal_rule",
                    "title": "資料共享管理辦法",
                    "source_system": "unit_test",
                    "language": "zh-TW",
                },
                strategy="local",
                output_dir=str(output_dir),
                no_db=True,
            )

            self.assertEqual(result["status"], "parsed_no_db")
            self.assertEqual(result["selected_strategy"], "local")
            self.assertGreaterEqual(result["total_chunks"], 1)
            self.assertTrue(result["summary_generated"])
            markdown_path = Path(result["markdown_path"])
            self.assertTrue(markdown_path.exists())
            markdown = markdown_path.read_text(encoding="utf-8")
            self.assertIn("# 資料共享管理辦法", markdown)
            self.assertIn("第一條", markdown)
            self.assertIn("資料共享", markdown)
            self.assertIn("parse_engine", result)
            self.assertIn("table_repair_status", result)

    def test_chunker_creates_parent_and_child_chunks(self) -> None:
        chunks = create_chunks(
            [
                StructuredSection(
                    section_index=1,
                    text="第一條 客戶資料共享應載明於契據文件。\n第二項 應保留查核紀錄。",
                    chunk_level="article",
                    source_structure_type="regulation_article",
                    heading_path="第一條",
                    section_title="第一條",
                    clause_number="第一條",
                )
            ]
        )

        self.assertGreaterEqual(len(chunks), 2)
        self.assertEqual(chunks[0].chunk_level, "parent")
        self.assertEqual(chunks[0].source_structure_type, "regulation_article")
        self.assertIsNone(chunks[0].parent_chunk_id)
        self.assertEqual(chunks[1].chunk_level, "child")
        self.assertEqual(chunks[1].parent_chunk_id, chunks[0].chunk_index)
        self.assertIn("第一條", chunks[1].chunk_text)

    def test_gemini_parent_schema_normalizes_parent_chunks(self) -> None:
        chunks = _chunk_records_from_json(
            {
                "parent_chunks": [
                    {
                        "chunk_index": 1,
                        "chunk_level": "article",
                        "heading_path": "第一條",
                        "section_title": "第一條",
                        "clause_number": "第一條",
                        "start_line": 2,
                        "end_line": 3,
                    }
                ]
            },
            markdown="# 客戶資料共享政策\n\n第一條 客戶資料共享應保留完整紀錄。",
        )

        self.assertEqual(len(chunks), 1)
        self.assertEqual(chunks[0].chunk_level, "parent")
        self.assertEqual(chunks[0].source_structure_type, "article")
        self.assertEqual(chunks[0].section_title, "第一條")
        self.assertEqual(chunks[0].chunk_text, "第一條 客戶資料共享應保留完整紀錄。")

    def test_local_chunks_can_use_gemini_parent_chunker(self) -> None:
        file_info = FileInfo(
            file_path=Path("policy.docx"),
            file_name="policy.docx",
            file_extension=".docx",
            file_size=100,
            file_type="docx",
        )
        metadata = prepare_metadata(
            file_info,
            {
                "file_path": str(file_info.file_path),
                "internal_code": "POLICY-LLM-PARENT",
                "document_type": "internal_rule",
                "title": "客戶資料共享政策",
            },
        )
        parent = ChunkRecord(
            chunk_index=1,
            chunk_level="parent",
            source_structure_type="regulation_article",
            heading_path="第一條",
            section_title="第一條",
            clause_number="第一條",
            page_start=None,
            page_end=None,
            chunk_text="第一條 客戶資料共享應保留完整紀錄。",
            token_count=20,
            char_count=20,
        )
        summary = SummaryResult(
            short_summary="測試摘要",
            keywords=["客戶資料"],
            main_topics=["資料共享"],
            summary_generated=True,
        )

        with tempfile.TemporaryDirectory() as temp_dir:
            markdown_path = Path(temp_dir) / "policy.md"
            markdown_path.write_text("# 客戶資料共享政策\n\n第一條 客戶資料共享應保留完整紀錄。", encoding="utf-8")
            with (
                patch("ingestion.hybrid.chunkers.read_local_document") as reader,
                patch("ingestion.hybrid.chunkers.chunk_markdown_with_gemini") as gemini_chunker,
            ):
                reader.return_value = SimpleNamespace(
                    blocks=[],
                    markdown=markdown_path.read_text(encoding="utf-8"),
                    output_markdown_path=markdown_path,
                    parse_engine="unit_test",
                    table_repair_status="not_needed",
                )
                gemini_chunker.return_value = (
                    [parent],
                    summary,
                    Path(temp_dir) / "chunks.json",
                    GeminiTokenUsage(
                        call_name="gemini_chunking",
                        model="test-model",
                        input_tokens=10,
                        output_tokens=5,
                    ),
                )

                prepared = prepare_local_chunks(
                    file_info,
                    metadata,
                    Path(temp_dir),
                    parent_chunker="gemini",
                )

        self.assertEqual(prepared.stage, "llm_parent_chunking")
        self.assertEqual(prepared.metrics["parent_chunker"], "gemini")
        self.assertEqual(prepared.chunks[0].chunk_level, "parent")
        self.assertEqual(prepared.chunks[1].chunk_level, "child")
        self.assertEqual(prepared.chunks[1].parent_chunk_id, prepared.chunks[0].chunk_index)

    def test_broken_table_detector_flags_misaligned_tables(self) -> None:
        good_table = """
| 項目 | 要求 |
| --- | --- |
| 資料共享 | 應保留紀錄 |
"""
        broken_table = """
| 類別 | 項目 | 說明 |
| --- | --- | --- |
| 補助 | A | 正常 |
|  |  | 合併儲存格造成空欄 |
|  |  | 又一列空欄 |
"""

        self.assertFalse(detect_broken_table(good_table))
        self.assertTrue(detect_broken_table(broken_table))

    def test_table_repair_skips_without_gemini_key(self) -> None:
        broken_table = """
| 類別 | 項目 | 說明 |
| --- | --- | --- |
| 補助 | A | 正常 |
|  |  | 合併儲存格造成空欄 |
|  |  | 又一列空欄 |
"""

        with patch("ingestion.advanced_markdown._gemini_key_available", return_value=False):
            repaired, status = repair_complex_tables(broken_table)

        self.assertEqual(repaired, broken_table)
        self.assertEqual(status, "skipped_no_gemini_key")

    def test_metadata_enrichment_extracts_version_stamps(self) -> None:
        file_info = FileInfo(
            file_path=Path("洗錢防制辦法2026年版.docx"),
            file_name="洗錢防制辦法2026年版.docx",
            file_extension=".docx",
            file_size=100,
            file_type="docx",
        )
        metadata = prepare_metadata(
            file_info,
            {
                "file_path": str(file_info.file_path),
                "internal_code": "AML-2026",
                "document_type": "internal_rule",
                "title": "洗錢防制辦法2026年版",
            },
        )
        enriched = enrich_metadata_from_markdown(
            metadata,
            """
# 洗錢防制辦法2026年版

發文單位：風控部
生效日期：2026年1月1日
修正日期：2025年12月20日
""",
        )

        self.assertEqual(enriched.issuing_unit, "風控部")
        self.assertEqual(enriched.responsible_unit, "風控部")
        self.assertEqual(enriched.effective_year, 2026)
        self.assertEqual(enriched.effective_date.isoformat(), "2026-01-01")
        self.assertEqual(enriched.revision_date.isoformat(), "2025-12-20")
        self.assertEqual(enriched.document_family, "洗錢防制辦法")
        self.assertEqual(enriched.normalized_version_label, "2026-01-01")

    def test_public_legacy_wrappers_reexport_legacy_implementations(self) -> None:
        self.assertIs(run_ingestion, legacy_local_ingestion)
        self.assertTrue(callable(legacy_gemini_ingestion))

    def test_folder_ingestion_uses_hybrid_pipeline(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "hybrid_policy.docx"
            docx_path.write_bytes(b"placeholder")

            with patch("ingestion.run_folder.run_hybrid_ingestion") as runner:
                runner.return_value = {
                    "status": "parsed_no_db",
                    "stage": "local_chunking",
                    "selected_strategy": "local",
                    "total_chunks": 3,
                }
                results = run_folder_ingestion(
                    temp_dir,
                    source_system="unit_test",
                    language="zh-TW",
                    internal_code_prefix="HYBRID",
                    strategy="auto",
                    no_db=True,
                )

        self.assertEqual(len(results), 1)
        self.assertEqual(results[0].selected_strategy, "local")
        self.assertEqual(results[0].total_chunks, 3)
        runner.assert_called_once()


if __name__ == "__main__":
    unittest.main()
