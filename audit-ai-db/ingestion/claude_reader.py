from __future__ import annotations

import base64
import os
import re
import time
from dataclasses import dataclass
from pathlib import Path

from ingestion.config import _load_dotenv_file
from ingestion.models import IngestionError


TEXT_HIGH = 200
TEXT_LOW = 50
PDF_DPI = 150
VISION_MODES = {"minimal", "full", "off"}


@dataclass(frozen=True)
class PageAnalysis:
    page: int
    text_chars: int
    image_count: int
    route: str
    text_preview: str


@dataclass(frozen=True)
class ClaudeTokenUsage:
    call_name: str
    model: str
    input_tokens: int
    output_tokens: int
    cache_creation_input_tokens: int = 0
    cache_read_input_tokens: int = 0

    @property
    def total_tokens(self) -> int:
        return self.input_tokens + self.output_tokens

    def to_dict(self) -> dict[str, int | str]:
        return {
            "call_name": self.call_name,
            "model": self.model,
            "input_tokens": self.input_tokens,
            "output_tokens": self.output_tokens,
            "cache_creation_input_tokens": self.cache_creation_input_tokens,
            "cache_read_input_tokens": self.cache_read_input_tokens,
            "total_tokens": self.total_tokens,
        }

    @classmethod
    def zero(cls, call_name: str = "none", model: str = "") -> "ClaudeTokenUsage":
        return cls(call_name=call_name, model=model, input_tokens=0, output_tokens=0)

    @classmethod
    def combine(
        cls,
        call_name: str,
        usages: list["ClaudeTokenUsage"],
    ) -> "ClaudeTokenUsage":
        model = usages[-1].model if usages else ""
        return cls(
            call_name=call_name,
            model=model,
            input_tokens=sum(usage.input_tokens for usage in usages),
            output_tokens=sum(usage.output_tokens for usage in usages),
            cache_creation_input_tokens=sum(
                usage.cache_creation_input_tokens for usage in usages
            ),
            cache_read_input_tokens=sum(
                usage.cache_read_input_tokens for usage in usages
            ),
        )


@dataclass(frozen=True)
class ClaudeCallResult:
    text: str
    usage: ClaudeTokenUsage


@dataclass(frozen=True)
class ClaudeReadResult:
    markdown: str
    page_analysis: list[PageAnalysis]
    output_markdown_path: Path | None = None
    usage: list[ClaudeTokenUsage] | None = None


def load_anthropic_settings() -> tuple[str, str]:
    project_root = Path(__file__).resolve().parents[1]
    _load_dotenv_file(project_root / ".env")

    api_key = os.getenv("ANTHROPIC_API_KEY")
    model = os.getenv("ANTHROPIC_MODEL", "claude-sonnet-4-5-20250929")
    if not api_key:
        raise IngestionError(
            "claude_configuration",
            "ANTHROPIC_API_KEY is required in .env for Claude ingestion",
        )
    return api_key, model


def create_anthropic_client(api_key: str):
    try:
        import anthropic
    except ImportError as exc:
        raise IngestionError(
            "claude_configuration",
            "anthropic is required. Install dependencies with: pip install -r requirements.txt",
        ) from exc

    return anthropic.Anthropic(api_key=api_key)


def route_page(text_chars: int, image_count: int, vision_mode: str = "minimal") -> str:
    if vision_mode not in VISION_MODES:
        raise IngestionError(
            "claude_configuration",
            f"vision_mode must be one of {sorted(VISION_MODES)}, got: {vision_mode}",
        )

    if vision_mode == "off":
        return "text"

    if text_chars >= TEXT_HIGH:
        return "text"
    if text_chars < TEXT_LOW and image_count > 0:
        return "claude_vision"
    if vision_mode == "full" and image_count > 0:
        return "mixed"
    return "text"


def analyze_pdf_pages(pdf_path: Path, vision_mode: str = "minimal") -> list[PageAnalysis]:
    try:
        import fitz
    except ImportError as exc:
        raise IngestionError(
            "claude_page_analysis",
            "PyMuPDF is required for Claude PDF page analysis",
        ) from exc

    pages: list[PageAnalysis] = []
    try:
        document = fitz.open(str(pdf_path))
    except Exception as exc:
        raise IngestionError("claude_page_analysis", f"unable to open PDF: {exc}") from exc

    try:
        for page_index in range(document.page_count):
            page = document.load_page(page_index)
            text = page.get_text("text").strip()
            image_count = len(page.get_images(full=True))
            route = route_page(len(text), image_count, vision_mode)
            pages.append(
                PageAnalysis(
                    page=page_index + 1,
                    text_chars=len(text),
                    image_count=image_count,
                    route=route,
                    text_preview=re.sub(r"\s+", " ", text[:120]),
                )
            )
    finally:
        document.close()

    return pages


def read_document_with_claude(
    file_path: Path,
    file_type: str,
    output_dir: Path,
    max_pages: int | None = None,
    vision_mode: str = "minimal",
    max_vision_pages: int | None = None,
    output_name: str | None = None,
) -> ClaudeReadResult:
    api_key, model = load_anthropic_settings()
    client = create_anthropic_client(api_key)

    output_dir.mkdir(parents=True, exist_ok=True)
    if file_type == "pdf":
        return _read_pdf_with_claude(
            client=client,
            model=model,
            pdf_path=file_path,
            output_dir=output_dir,
            max_pages=max_pages,
            vision_mode=vision_mode,
            max_vision_pages=max_vision_pages,
            output_name=output_name,
        )
    if file_type == "docx":
        return _read_docx_as_markdown(file_path, output_dir, output_name=output_name)
    raise IngestionError(
        "claude_reading",
        f"Claude ingestion currently supports .pdf and .docx, got: {file_type}",
    )


def _read_pdf_with_claude(
    *,
    client,
    model: str,
    pdf_path: Path,
    output_dir: Path,
    max_pages: int | None,
    vision_mode: str,
    max_vision_pages: int | None,
    output_name: str | None,
) -> ClaudeReadResult:
    pages = analyze_pdf_pages(pdf_path, vision_mode)
    selected_pages = pages[:max_pages] if max_pages else pages
    if not selected_pages:
        raise IngestionError("claude_reading", "PDF has no pages")

    vision_pages = [
        page.page for page in selected_pages if page.route in {"claude_vision", "mixed"}
    ]
    if max_vision_pages is not None and len(vision_pages) > max_vision_pages:
        raise IngestionError(
            "claude_reading",
            (
                "vision page budget exceeded: "
                f"{len(vision_pages)} pages need Claude Vision, limit is {max_vision_pages}"
            ),
            status="skipped_vision_budget",
        )

    safe_output_name = _safe_output_name(output_name or pdf_path.stem)
    parts: list[str] = [f"# {pdf_path.stem}", ""]
    usage: list[ClaudeTokenUsage] = []
    for page_info in selected_pages:
        parts.append(f"\n\n<!-- page:{page_info.page} route:{page_info.route} -->\n")
        if page_info.route == "text":
            text = _extract_pdf_page_text(pdf_path, page_info.page)
            parts.append(f"## Page {page_info.page}\n\n{text}")
            continue

        image_path = _render_pdf_page(
            pdf_path,
            page_info.page,
            output_dir / "page_images",
            output_stem=safe_output_name,
        )
        prompt = _page_markdown_prompt(page_info.page)
        vision_result = _call_claude_image(
            client=client,
            model=model,
            image_path=image_path,
            prompt=prompt,
            max_tokens=4096,
            call_name=f"claude_vision_page_{page_info.page}",
        )
        usage.append(vision_result.usage)

        if page_info.route == "mixed":
            extracted_text = _extract_pdf_page_text(pdf_path, page_info.page)
            parts.append(f"## Page {page_info.page} Extracted Text\n\n{extracted_text}")
            parts.append(f"\n\n## Page {page_info.page} Claude Vision Supplement\n\n{vision_result.text}")
        else:
            parts.append(f"## Page {page_info.page}\n\n{vision_result.text}")

    markdown = "\n".join(parts).strip() + "\n"
    markdown_path = output_dir / "markdown" / f"{safe_output_name}.md"
    markdown_path.parent.mkdir(parents=True, exist_ok=True)
    markdown_path.write_text(markdown, encoding="utf-8")
    return ClaudeReadResult(
        markdown=markdown,
        page_analysis=selected_pages,
        output_markdown_path=markdown_path,
        usage=usage,
    )


def _read_docx_as_markdown(
    file_path: Path,
    output_dir: Path,
    *,
    output_name: str | None,
) -> ClaudeReadResult:
    try:
        from docx import Document
    except ImportError as exc:
        raise IngestionError(
            "claude_reading",
            "python-docx is required for DOCX Claude ingestion",
        ) from exc

    try:
        document = Document(str(file_path))
    except Exception as exc:
        raise IngestionError("claude_reading", f"unable to open DOCX: {exc}") from exc

    lines = [f"# {file_path.stem}", ""]
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name.lower() if paragraph.style else ""
        if style.startswith("heading"):
            level = "##" if "1" in style else "###"
            lines.append(f"{level} {text}")
        else:
            lines.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        lines.append(f"\n## Table {table_index}\n")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                lines.append("| " + " | ".join(cells) + " |")

    markdown = "\n\n".join(lines).strip() + "\n"
    safe_output_name = _safe_output_name(output_name or file_path.stem)
    markdown_path = output_dir / "markdown" / f"{safe_output_name}.md"
    markdown_path.parent.mkdir(parents=True, exist_ok=True)
    markdown_path.write_text(markdown, encoding="utf-8")
    return ClaudeReadResult(
        markdown=markdown,
        page_analysis=[],
        output_markdown_path=markdown_path,
        usage=[],
    )


def _extract_pdf_page_text(pdf_path: Path, page_number: int) -> str:
    import fitz

    document = fitz.open(str(pdf_path))
    try:
        page = document.load_page(page_number - 1)
        return page.get_text("text").strip()
    finally:
        document.close()


def _render_pdf_page(
    pdf_path: Path,
    page_number: int,
    image_dir: Path,
    *,
    output_stem: str | None = None,
) -> Path:
    import fitz

    image_dir.mkdir(parents=True, exist_ok=True)
    document = fitz.open(str(pdf_path))
    try:
        page = document.load_page(page_number - 1)
        pixmap = page.get_pixmap(matrix=fitz.Matrix(PDF_DPI / 72, PDF_DPI / 72), alpha=False)
        stem = output_stem or _safe_output_name(pdf_path.stem)
        output_path = image_dir / f"{stem}_page_{page_number:03d}.png"
        pixmap.save(str(output_path))
        return output_path
    finally:
        document.close()


def _call_claude_image(
    *,
    client,
    model: str,
    image_path: Path,
    prompt: str,
    max_tokens: int,
    call_name: str,
) -> ClaudeCallResult:
    image_b64 = base64.b64encode(image_path.read_bytes()).decode("ascii")
    for attempt in range(3):
        try:
            message = client.messages.create(
                model=model,
                max_tokens=max_tokens,
                system=(
                    "You convert internal business/legal document page images into faithful Markdown. "
                    "Do not summarize. Do not invent text. Preserve visible headings, lists, tables, and legal markers."
                ),
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image",
                                "source": {
                                    "type": "base64",
                                    "media_type": "image/png",
                                    "data": image_b64,
                                },
                            },
                            {"type": "text", "text": prompt},
                        ],
                    }
                ],
            )
            return ClaudeCallResult(
                text=_message_text(message).strip(),
                usage=_message_usage(message, call_name, model),
            )
        except Exception as exc:
            if attempt == 2:
                raise IngestionError("claude_reading", f"Claude vision failed: {exc}") from exc
            time.sleep(2**attempt)
    return ClaudeCallResult(
        text="",
        usage=ClaudeTokenUsage(call_name=call_name, model=model, input_tokens=0, output_tokens=0),
    )


def call_claude_text(
    *,
    prompt: str,
    system: str,
    max_tokens: int = 8192,
    call_name: str = "claude_text",
) -> ClaudeCallResult:
    api_key, model = load_anthropic_settings()
    client = create_anthropic_client(api_key)
    for attempt in range(3):
        try:
            message = client.messages.create(
                model=model,
                max_tokens=max_tokens,
                system=system,
                messages=[{"role": "user", "content": prompt}],
            )
            return ClaudeCallResult(
                text=_message_text(message).strip(),
                usage=_message_usage(message, call_name, model),
            )
        except Exception as exc:
            if attempt == 2:
                raise IngestionError("claude_chunking", f"Claude text call failed: {exc}") from exc
            time.sleep(2**attempt)
    return ClaudeCallResult(
        text="",
        usage=ClaudeTokenUsage(call_name=call_name, model=model, input_tokens=0, output_tokens=0),
    )


def _message_text(message) -> str:
    parts = []
    for block in getattr(message, "content", []):
        if getattr(block, "type", None) == "text":
            parts.append(getattr(block, "text", ""))
    return "\n".join(parts)


def _message_usage(message, call_name: str, model: str) -> ClaudeTokenUsage:
    usage = getattr(message, "usage", None)
    return ClaudeTokenUsage(
        call_name=call_name,
        model=model,
        input_tokens=_usage_int(usage, "input_tokens"),
        output_tokens=_usage_int(usage, "output_tokens"),
        cache_creation_input_tokens=_usage_int(usage, "cache_creation_input_tokens"),
        cache_read_input_tokens=_usage_int(usage, "cache_read_input_tokens"),
    )


def _usage_int(usage, field_name: str) -> int:
    if usage is None:
        return 0
    value = getattr(usage, field_name, 0)
    try:
        return int(value or 0)
    except (TypeError, ValueError):
        return 0


def _safe_output_name(value: str) -> str:
    safe = re.sub(r'[\\/:"*?<>|]+', "_", value).strip()
    return safe or "document"


def _page_markdown_prompt(page_number: int) -> str:
    return (
        f"Read page {page_number} and convert all visible document content into Markdown.\n"
        "Requirements:\n"
        "- Preserve original Chinese text as faithfully as possible.\n"
        "- Preserve legal/regulatory markers such as 第一條, 第二條, 一、, 二、, 法務室意見, 結論.\n"
        "- Preserve tables as Markdown tables when visible.\n"
        "- If text is unclear, mark it as [辨識不清].\n"
        "- Do not add explanations outside the Markdown content."
    )
